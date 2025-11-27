from __future__ import annotations

import io
import json
from typing import Iterable, List, Sequence, Tuple

import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.database.sqlite_dal import QuizQuestion, QuizTopic


OPTION_PREFIX_PATTERN = re.compile(r"^[A-Za-z]\s*[\.\)\-:]\s+")


def _clean_option_text(option: str) -> str:
    stripped = option.strip()
    if OPTION_PREFIX_PATTERN.match(stripped):
        stripped = OPTION_PREFIX_PATTERN.sub("", stripped, count=1).strip()
    return stripped


def _normalize_options(options: Sequence[str] | str | None) -> List[str]:
    if options is None:
        return []
    if isinstance(options, list):
        return [_clean_option_text(str(option)) for option in options]
    if isinstance(options, str):
        try:
            parsed = json.loads(options)
            if isinstance(parsed, list):
                return [_clean_option_text(str(option)) for option in parsed]
        except json.JSONDecodeError:
            pass
        return [_clean_option_text(options)]
    return [_clean_option_text(str(options))]


def _resolve_correct_answer(question: QuizQuestion, options: List[str]) -> Tuple[str, str]:
    raw_value = question.right_option
    option_index = None

    if raw_value is None:
        option_index = None
    else:
        if isinstance(raw_value, int):
            option_index = raw_value
        elif isinstance(raw_value, str):
            raw_value_stripped = raw_value.strip()
            if raw_value_stripped.isdigit():
                option_index = int(raw_value_stripped)
            elif len(raw_value_stripped) == 1 and raw_value_stripped.isalpha():
                option_index = ord(raw_value_stripped.upper()) - 65
            else:
                try:
                    option_index = options.index(raw_value_stripped)
                except ValueError:
                    option_index = None

    if option_index is not None and 0 <= option_index < len(options):
        option_label = chr(65 + option_index)
        option_text = options[option_index]
    else:
        option_label = str(raw_value) if raw_value is not None else "N/A"
        option_text = "N/A"

    return option_label, option_text


def build_quiz_docx(topic: QuizTopic, questions: Iterable[QuizQuestion]) -> bytes:
    question_list = list(questions)
    document = Document()
    document.add_heading(f"{topic.topic} Quiz", level=1)

    meta_paragraph = document.add_paragraph()
    meta_paragraph.add_run(f"Category: {topic.category} • {topic.subcategory}").italic = True
    if topic.difficulty:
        meta_paragraph.add_run(f" • Difficulty: {topic.difficulty.title()}")

    document.add_paragraph()

    for idx, question in enumerate(question_list, start=1):
        document.add_paragraph(f"{idx}. {question.question}", style="List Number")
        options = _normalize_options(question.options)
        for opt_idx, option in enumerate(options):
            option_run = document.add_paragraph(style="List Bullet").add_run(
                f"{chr(65 + opt_idx)}. {option}"
            )
            option_run.font.size = Pt(11)
        document.add_paragraph()

    document.add_page_break()
    document.add_heading("Answer Key", level=2)
    for idx, question in enumerate(question_list, start=1):
        options = _normalize_options(question.options)
        label, text = _resolve_correct_answer(question, options)
        document.add_paragraph(f"{idx}. {label} - {text}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_quiz_pdf(topic: QuizTopic, questions: Iterable[QuizQuestion]) -> bytes:
    question_list = list(questions)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )

    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    heading_style = styles["Heading2"]
    normal = ParagraphStyle(
        "QuizNormal",
        parent=styles["Normal"],
        leading=16,
        spaceAfter=6,
    )

    story = [
        Paragraph(f"{topic.topic} Quiz", title_style),
        Spacer(1, 12),
        Paragraph(
            f"Category: {topic.category} • {topic.subcategory}"
            + (f" • Difficulty: {topic.difficulty.title()}" if topic.difficulty else ""),
            normal,
        ),
        Spacer(1, 18),
    ]

    for idx, question in enumerate(question_list, start=1):
        story.append(Paragraph(f"{idx}. {question.question}", heading_style))
        options = _normalize_options(question.options)
        for opt_idx, option in enumerate(options):
            story.append(
                Paragraph(f"{chr(65 + opt_idx)}. {option}", normal)
            )
        story.append(Spacer(1, 12))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Answer Key", heading_style))

    for idx, question in enumerate(question_list, start=1):
        options = _normalize_options(question.options)
        label, text = _resolve_correct_answer(question, options)
        story.append(Paragraph(f"{idx}. {label} - {text}", normal))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


